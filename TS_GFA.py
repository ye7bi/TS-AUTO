#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur automatique de Termsheet Word
Interface PyQt6 pour remplir un template de termsheet et générer un fichier Word final
"""

import sys
import os
import re
import pandas as pd
from PyQt6.QtWidgets import QDialog  # Ajouter QDialog aux imports existants
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QFormLayout, QLineEdit, QTextEdit, QCheckBox, QLabel, QPushButton,
    QFileDialog, QMessageBox, QScrollArea, QGroupBox, QSpinBox,
    QComboBox, QFrame, QDoubleSpinBox, QDialog  # <- Ajouter QDialog ici
)

from PyQt6.QtCore import Qt, pyqtSignal
from PyQt6.QtGui import QFont, QPixmap

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Erreur: Le module python-docx n'est pas installé.")
    print("Installez-le avec: pip install python-docx")
    sys.exit(1)


class NumberToWords:
    """Convertit les nombres en lettres (français)"""
    
    UNITS = ['', 'un', 'deux', 'trois', 'quatre', 'cinq', 'six', 'sept', 'huit', 'neuf']
    TEENS = ['dix', 'onze', 'douze', 'treize', 'quatorze', 'quinze', 'seize', 
             'dix-sept', 'dix-huit', 'dix-neuf']
    TENS = ['', '', 'vingt', 'trente', 'quarante', 'cinquante', 'soixante', 
            'soixante-dix', 'quatre-vingt', 'quatre-vingt-dix']
    
    @classmethod
    def convert(cls, number: int) -> str:
        """Convertit un nombre en lettres"""
        if number == 0:
            return "zéro"
        
        if number < 0:
            return "moins " + cls.convert(-number)
        
        result = []
        
        # Milliards
        if number >= 1000000000:
            billions = number // 1000000000
            if billions == 1:
                result.append("un milliard")
            else:
                result.append(cls._convert_hundreds(billions) + " milliards")
            number %= 1000000000
        
        # Millions
        if number >= 1000000:
            millions = number // 1000000
            if millions == 1:
                result.append("un million")
            else:
                result.append(cls._convert_hundreds(millions) + " millions")
            number %= 1000000
        
        # Milliers
        if number >= 1000:
            thousands = number // 1000
            if thousands == 1:
                result.append("mille")
            else:
                result.append(cls._convert_hundreds(thousands) + " mille")
            number %= 1000
        
        # Centaines
        if number > 0:
            result.append(cls._convert_hundreds(number))
        
        return " ".join(result)
    
    @classmethod
    def _convert_hundreds(cls, number: int) -> str:
        """Convertit un nombre de 0 à 999 en lettres"""
        result = []
        
        # Centaines
        if number >= 100:
            hundreds = number // 100
            if hundreds == 1:
                result.append("cent")
            else:
                result.append(cls.UNITS[hundreds] + " cent")
            number %= 100
        
        # Dizaines et unités
        if number >= 20:
            tens = number // 10
            units = number % 10
            if tens == 7:
                if units == 1:
                    result.append("soixante et onze")
                elif units > 1:
                    result.append("soixante-" + cls.TEENS[units - 10])
                else:
                    result.append("soixante-dix")
            elif tens == 9:
                if units > 0:
                    result.append("quatre-vingt-" + cls.TEENS[units - 10])
                else:
                    result.append("quatre-vingt-dix")
            else:
                tens_word = cls.TENS[tens]
                if units == 1 and tens != 8:
                    result.append(tens_word + " et un")
                elif units > 0:
                    result.append(tens_word + "-" + cls.UNITS[units])
                else:
                    if tens == 8:
                        result.append("quatre-vingts")
                    else:
                        result.append(tens_word)
        elif number >= 10:
            result.append(cls.TEENS[number - 10])
        elif number > 0:
            result.append(cls.UNITS[number])
        
        return " ".join(result)
    
def format_number_with_dots(number: str) -> str:
    """Formatte un nombre sous forme 1.000.000"""
    try:
        n = int(number.replace(" ", "").replace(",", ""))
        return f"{n:,}".replace(",", ".")
    except:
        return number  # Si non numérique, on retourne brut



class ClauseWidget(QWidget):
    """Widget pour une clause optionnelle avec ses champs associés"""
    
    def __init__(self, clause_name: str, clause_text: str, fields: List[Dict] = None):
        super().__init__()
        self.clause_name = clause_name
        self.clause_text = clause_text
        self.fields = fields or []
        self.field_widgets = {}
        
        self.setup_ui()
    
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Case à cocher pour activer/désactiver la clause
        self.checkbox = QCheckBox(self.clause_name)
        self.checkbox.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        self.checkbox.toggled.connect(self.toggle_fields)
        layout.addWidget(self.checkbox)
        
        # Container pour les champs
        self.fields_container = QWidget()
        fields_layout = QFormLayout()
        
        # Créer les champs
        for field in self.fields:
            field_name = field['name']
            field_type = field.get('type', 'text')
            
            if field_type == 'number':
                widget = QSpinBox()
                widget.setMaximum(999)
                widget.setMinimum(0)
            else:
                widget = QLineEdit()
            
            self.field_widgets[field_name] = widget
            fields_layout.addRow(field['label'], widget)
        
        self.fields_container.setLayout(fields_layout)
        self.fields_container.setVisible(False)  # Masqué par défaut
        layout.addWidget(self.fields_container)
        
        self.setLayout(layout)
    
    def toggle_fields(self, checked: bool):
        """Affiche/masque les champs selon l'état de la case à cocher"""
        self.fields_container.setVisible(checked)
    
    def is_enabled(self) -> bool:
        """Retourne True si la clause est activée"""
        return self.checkbox.isChecked()
    
    def get_field_values(self) -> Dict[str, str]:
        """Retourne les valeurs des champs"""
        values = {}
        for field_name, widget in self.field_widgets.items():
            if isinstance(widget, QSpinBox):
                values[field_name] = str(widget.value())
            else:
                values[field_name] = widget.text().strip()
        return values



class ProfileDialog(QDialog):
    """Boîte de dialogue pour créer/modifier un profil"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Nouveau Profil Promoteur")
        self.setModal(True)
        self.setFixedSize(400, 250)
        
        layout = QVBoxLayout()
        
        # Formulaire
        form_layout = QFormLayout()
        
        self.nom_promoteur_edit = QLineEdit()
        self.nom_contact_edit = QLineEdit()
        self.adresse_promoteur_edit = QLineEdit()
        
        self.civilite_combo = QComboBox()
        self.civilite_combo.addItems(['Monsieur', 'Madame', 'Messieurs'])
        
        form_layout.addRow("Nom du promoteur :", self.nom_promoteur_edit)
        form_layout.addRow("Nom du contact :", self.nom_contact_edit)
        form_layout.addRow("Adresse du promoteur :", self.adresse_promoteur_edit)
        form_layout.addRow("Civilité :", self.civilite_combo)
        
        layout.addLayout(form_layout)
        
        # Boutons
        buttons_layout = QHBoxLayout()
        
        self.ok_button = QPushButton("Enregistrer")
        self.ok_button.clicked.connect(self.accept)
        
        self.cancel_button = QPushButton("Annuler")
        self.cancel_button.clicked.connect(self.reject)
        
        buttons_layout.addWidget(self.cancel_button)
        buttons_layout.addWidget(self.ok_button)
        
        layout.addLayout(buttons_layout)
        self.setLayout(layout)
    
    def get_profile_data(self):
        """Récupère les données du profil"""
        return {
            'nom_promoteur': self.nom_promoteur_edit.text().strip(),
            'nom_contact': self.nom_contact_edit.text().strip(),
            'adresse_promoteur': self.adresse_promoteur_edit.text().strip(),
            'civilite': self.civilite_combo.currentText()
        }


class TermsheetGenerator(QMainWindow):
    """Application principale pour générer les termsheets"""
    
    def __init__(self):
        super().__init__()
        self.template_path = None
        self.profiles_path = Path("profils.xlsx") 
        self.setup_ui()
        self.setup_clauses()
        self.load_profiles()
        self.load_default_template()
    
    def load_default_template(self):
        """Charge automatiquement le template par défaut"""
        default_template = "template_ts.docx"
        
        if os.path.exists(default_template):
            self.template_path = default_template
            self.template_label.setText(f"Template: {default_template}")
            self.generate_button.setEnabled(True)
            self.preview_button.setEnabled(True)
        else:
            self.template_label.setText("Template par défaut non trouvé - Veuillez importer un template")
            self.generate_button.setEnabled(False)
            self.preview_button.setEnabled(False)
    
    def setup_ui(self):
        """Configure l'interface utilisateur"""
        self.setWindowTitle("Générateur de Termsheet - LCL")
        self.setGeometry(100, 100, 900, 800)
        
        # Widget principal avec scroll
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        main_layout = QVBoxLayout()
        
        # Zone de scroll
        scroll = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout()
        
        # Section import template
        template_group = QGroupBox("Template Word")
        template_layout = QHBoxLayout()
        
        self.template_label = QLabel("Chargement du template...")
        self.template_button = QPushButton("Changer de Template (.docx)")
        self.template_button.clicked.connect(self.import_template)
        
        template_layout.addWidget(self.template_label)
        template_layout.addWidget(self.template_button)
        template_group.setLayout(template_layout)
        
        scroll_layout.addWidget(template_group)

        # Section gestion des profils
        profils_group = QGroupBox("Gestion des Profils Promoteurs")
        profils_layout = QHBoxLayout()
        
        # Menu déroulant pour sélectionner un profil
        self.profil_combo = QComboBox()
        self.profil_combo.addItem("-- Sélectionner un profil --")
        self.profil_combo.currentTextChanged.connect(self.on_profil_selected)
        
        # Bouton pour créer un nouveau profil
        self.nouveau_profil_button = QPushButton("Nouveau Profil")
        self.nouveau_profil_button.clicked.connect(self.create_new_profile)
        
        profils_layout.addWidget(QLabel("Profil :"))
        profils_layout.addWidget(self.profil_combo)
        profils_layout.addWidget(self.nouveau_profil_button)
        profils_layout.addStretch()
        
        profils_group.setLayout(profils_layout)
        scroll_layout.addWidget(profils_group)
        
        # Section informations principales
        main_info_group = QGroupBox("Informations Principales")
        main_info_layout = QFormLayout()
        
        # Créer tous les champs principaux
        self.fields = {}
        
        main_fields = [
            ('nom_promoteur', 'Nom du promoteur'),
            ('nom_contact', 'Nom du contact'),
            ('adresse_promoteur', 'Adresse du promoteur'),
            ('date', 'Date'),
            ('ville', 'Ville'),
            ('reference_dossier', 'Référence dossier'),
            ('nom_sccv', 'Nom de la SCCV (Emprunteur)'),
            ('numero_siren', 'Numéro SIREN'),
            ('ville_rcs', 'Ville RCS'),
            ('montant_credit', 'Montant du crédit promoteur (€)'),
            ('montant_gfa', 'Prix de (€ HT)'),
            ('frais_dossier', 'Frais de dossier (€)'),
            ('montant_apports', 'Montant des apports investis'),
            ('date_echeance_gfa', 'Date d\'échéance de la GFA')
        ]
        
        for field_key, field_label in main_fields:
            if field_key in ['montant_credit', 'montant_gfa', 'frais_dossier', 'montant_apports']:
                # Champs numériques pour les montants
                widget = QLineEdit()
                widget.setPlaceholderText("Entrez le montant en euros")
            else:
                widget = QLineEdit()
            
            self.fields[field_key] = widget
            main_info_layout.addRow(field_label, widget)
        
        # Civilité (dropdown)
        self.civilite_combo = QComboBox()
        self.civilite_combo.addItems(['Monsieur', 'Madame', 'Messieurs'])
        main_info_layout.addRow('Civilité', self.civilite_combo)
        
        # Objet/Description (zone de texte plus grande)
        self.objet_text = QTextEdit()
        self.objet_text.setMaximumHeight(80)
        self.objet_text.setPlainText("Réalisation à,…. , d'un immeuble neuf d'une surface de plancher de m² conçu en R+ comprenant logements et places de stationnement")
        main_info_layout.addRow('Objet / Description du programme', self.objet_text)
        
        main_info_group.setLayout(main_info_layout)
        scroll_layout.addWidget(main_info_group)
        
        # Section Taux et Conditions
        taux_group = QGroupBox("Taux et Conditions")
        taux_layout = QFormLayout()
        
        # Cases à cocher pour activer/désactiver les sections
        self.conditions_speculatives_checkbox = QCheckBox("Inclure les conditions spéculatives")
        self.conditions_speculatives_checkbox.setChecked(True)
        taux_layout.addRow(self.conditions_speculatives_checkbox)
        
        # Taux spéculatif
        self.taux_speculatif = QDoubleSpinBox()
        self.taux_speculatif.setDecimals(2)
        self.taux_speculatif.setMaximum(99.99)
        self.taux_speculatif.setValue(2.25)
        self.taux_speculatif.setSuffix(' %')
        taux_layout.addRow('Taux spéculatif', self.taux_speculatif)
        
        # Taux commission engagement spéculatif
        self.taux_comission_engagement_speculatif = QDoubleSpinBox()
        self.taux_comission_engagement_speculatif.setDecimals(2)
        self.taux_comission_engagement_speculatif.setMaximum(99.99)
        self.taux_comission_engagement_speculatif.setValue(0.75)
        self.taux_comission_engagement_speculatif.setSuffix(' %')
        taux_layout.addRow('Taux commission engagement spéculatif', self.taux_comission_engagement_speculatif)
        
        # Séparateur
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.Shape.HLine)
        separator1.setFrameShadow(QFrame.Shadow.Sunken)
        taux_layout.addRow(separator1)
        
        # Cases à cocher pour les conditions non spéculatives
        self.conditions_non_speculatives_checkbox = QCheckBox("Inclure les conditions non spéculatives")
        self.conditions_non_speculatives_checkbox.setChecked(True)
        taux_layout.addRow(self.conditions_non_speculatives_checkbox)
        
        # Taux non spéculatif
        self.taux_non_speculatif = QDoubleSpinBox()
        self.taux_non_speculatif.setDecimals(2)
        self.taux_non_speculatif.setMaximum(99.99)
        self.taux_non_speculatif.setValue(1.50)
        self.taux_non_speculatif.setSuffix(' %')
        taux_layout.addRow('Taux non spéculatif', self.taux_non_speculatif)
        
        # Taux commission engagement non spéculatif
        self.taux_comission_engagement_non_speculatif = QDoubleSpinBox()
        self.taux_comission_engagement_non_speculatif.setDecimals(2)
        self.taux_comission_engagement_non_speculatif.setMaximum(99.99)
        self.taux_comission_engagement_non_speculatif.setValue(0.50)
        self.taux_comission_engagement_non_speculatif.setSuffix(' %')
        taux_layout.addRow('Taux commission engagement non spéculatif', self.taux_comission_engagement_non_speculatif)
        
        # Séparateur
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.Shape.HLine)
        separator2.setFrameShadow(QFrame.Shadow.Sunken)
        taux_layout.addRow(separator2)
        
        # Taux commission forfaitaire
        self.taux_comission_forfaitaire = QDoubleSpinBox()
        self.taux_comission_forfaitaire.setDecimals(2)
        self.taux_comission_forfaitaire.setMaximum(99.99)
        self.taux_comission_forfaitaire.setValue(0.55)
        self.taux_comission_forfaitaire.setSuffix(' %')
        taux_layout.addRow('Taux commission forfaitaire', self.taux_comission_forfaitaire)
        
        taux_group.setLayout(taux_layout)
        scroll_layout.addWidget(taux_group)
        
        # Section Commercialisation
        commercialisation_group = QGroupBox("Niveaux de Commercialisation")
        commercialisation_layout = QFormLayout()
        
        # Layout horizontal pour niveau global + case à cocher apports
        niveau_global_layout = QHBoxLayout()
        self.niveau_commercialisation = QDoubleSpinBox()
        self.niveau_commercialisation.setDecimals(0)
        self.niveau_commercialisation.setMaximum(100)
        self.niveau_commercialisation.setValue(50)
        self.niveau_commercialisation.setSuffix(' %')
        
        self.inclure_apports_checkbox = QCheckBox("(en y ajoutant les apports)")
        self.inclure_apports_checkbox.setChecked(True)
        
        niveau_global_layout.addWidget(self.niveau_commercialisation)
        niveau_global_layout.addWidget(self.inclure_apports_checkbox)
        niveau_global_layout.addStretch()
        
        niveau_global_widget = QWidget()
        niveau_global_widget.setLayout(niveau_global_layout)
        commercialisation_layout.addRow('Niveau commercialisation global', niveau_global_widget)
        
        commercialisation_group.setLayout(commercialisation_layout)
        scroll_layout.addWidget(commercialisation_group)
        
        # Section clauses optionnelles
        self.clauses_group = QGroupBox("Clauses Optionnelles")
        self.clauses_layout = QVBoxLayout()
        self.clauses_group.setLayout(self.clauses_layout)
        scroll_layout.addWidget(self.clauses_group)
        
        # Boutons d'action
        buttons_layout = QHBoxLayout()
        
        self.generate_button = QPushButton("Générer le Termsheet")
        self.generate_button.clicked.connect(self.generate_termsheet)
        self.generate_button.setEnabled(False)
        
        self.preview_button = QPushButton("Aperçu")
        self.preview_button.clicked.connect(self.preview_termsheet)
        self.preview_button.setEnabled(False)
        
        buttons_layout.addWidget(self.preview_button)
        buttons_layout.addWidget(self.generate_button)
        
        scroll_layout.addLayout(buttons_layout)
        
        scroll_widget.setLayout(scroll_layout)
        scroll.setWidget(scroll_widget)
        scroll.setWidgetResizable(True)
        
        main_layout.addWidget(scroll)
        main_widget.setLayout(main_layout)


    
    def load_profiles(self):
        """Charge les profils depuis le fichier Excel"""
        try:
            if self.profiles_path.exists():
                df = pd.read_excel(self.profiles_path)
                
                # Vérifier que les colonnes existent
                required_columns = ['Nom du promoteur', 'Nom du contact', 'Adresse du promoteur', 'Civilité']
                if all(col in df.columns for col in required_columns):
                    # Ajouter les profils au combo
                    for promoteur in df['Nom du promoteur'].dropna():
                        self.profil_combo.addItem(str(promoteur))
                else:
                    print("Le fichier Excel ne contient pas les bonnes colonnes")
            else:
                # Créer le fichier Excel avec les en-têtes
                self.create_empty_profiles_file()
                
        except Exception as e:
            print(f"Erreur lors du chargement des profils : {e}")

    def create_empty_profiles_file(self):
        """Crée un fichier Excel vide avec les en-têtes"""
        try:
            df = pd.DataFrame(columns=['Nom du promoteur', 'Nom du contact', 'Adresse du promoteur', 'Civilité'])
            df.to_excel(self.profiles_path, index=False)
        except Exception as e:
            print(f"Erreur lors de la création du fichier profils : {e}")

    def on_profil_selected(self, profil_name):
        """Méthode appelée quand un profil est sélectionné"""
        if profil_name == "-- Sélectionner un profil --" or not profil_name:
            return
            
        try:
            if self.profiles_path.exists():
                df = pd.read_excel(self.profiles_path)
                
                # Trouver la ligne correspondant au profil
                profil_row = df[df['Nom du promoteur'] == profil_name]
                
                if not profil_row.empty:
                    # Remplir les champs
                    self.fields['nom_promoteur'].setText(str(profil_row.iloc[0]['Nom du promoteur']))
                    self.fields['nom_contact'].setText(str(profil_row.iloc[0]['Nom du contact']))
                    self.fields['adresse_promoteur'].setText(str(profil_row.iloc[0]['Adresse du promoteur']))
                    
                    # Sélectionner la civilité dans le combo
                    civilite = str(profil_row.iloc[0]['Civilité'])
                    if civilite in ['Monsieur', 'Madame', 'Messieurs']:
                        self.civilite_combo.setCurrentText(civilite)
                        
        except Exception as e:
            QMessageBox.warning(self, "Erreur", f"Erreur lors du chargement du profil : {e}")

    def create_new_profile(self):
        """Ouvre une boîte de dialogue pour créer un nouveau profil"""
        dialog = ProfileDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            profile_data = dialog.get_profile_data()
            self.save_profile_to_excel(profile_data)
            
            # Rafraîchir la liste des profils
            self.profil_combo.clear()
            self.profil_combo.addItem("-- Sélectionner un profil --")
            self.load_profiles()
            
            # Sélectionner le nouveau profil
            self.profil_combo.setCurrentText(profile_data['nom_promoteur'])

    def save_profile_to_excel(self, profile_data):
        """Sauvegarde un profil dans le fichier Excel"""
        try:
            # Charger le fichier existant ou créer un nouveau DataFrame
            if self.profiles_path.exists():
                df = pd.read_excel(self.profiles_path)
            else:
                df = pd.DataFrame(columns=['Nom du promoteur', 'Nom du contact', 'Adresse du promoteur', 'Civilité'])
            
            # Vérifier si le profil existe déjà
            if profile_data['nom_promoteur'] in df['Nom du promoteur'].values:
                # Mettre à jour
                df.loc[df['Nom du promoteur'] == profile_data['nom_promoteur']] = [
                    profile_data['nom_promoteur'],
                    profile_data['nom_contact'],
                    profile_data['adresse_promoteur'],
                    profile_data['civilite']
                ]
            else:
                # Ajouter nouvelle ligne
                new_row = pd.DataFrame([{
                    'Nom du promoteur': profile_data['nom_promoteur'],
                    'Nom du contact': profile_data['nom_contact'],
                    'Adresse du promoteur': profile_data['adresse_promoteur'],
                    'Civilité': profile_data['civilite']
                }])
                df = pd.concat([df, new_row], ignore_index=True)
            
            # Sauvegarder
            df.to_excel(self.profiles_path, index=False)
            QMessageBox.information(self, "Succès", "Profil sauvegardé avec succès !")
            
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la sauvegarde : {e}")

    def check_and_propose_save_profile(self):
        """Vérifie si un profil doit être proposé à la sauvegarde"""
        # Vérifier si aucun profil n'est sélectionné
        current_profil = self.profil_combo.currentText()
        if current_profil == "-- Sélectionner un profil --":
            
            # Vérifier si les 4 champs sont remplis
            nom_promoteur = self.fields['nom_promoteur'].text().strip()
            nom_contact = self.fields['nom_contact'].text().strip()
            adresse_promoteur = self.fields['adresse_promoteur'].text().strip()
            civilite = self.civilite_combo.currentText()
            
            if nom_promoteur and nom_contact and adresse_promoteur and civilite:
                # Proposer de sauvegarder
                reply = QMessageBox.question(
                    self,
                    "Enregistrer le promoteur",
                    "Voulez-vous enregistrer ce promoteur dans vos profils ?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                
                if reply == QMessageBox.StandardButton.Yes:
                    profile_data = {
                        'nom_promoteur': nom_promoteur,
                        'nom_contact': nom_contact,
                        'adresse_promoteur': adresse_promoteur,
                        'civilite': civilite
                    }
                    self.save_profile_to_excel(profile_data)
                    
                    # Rafraîchir la liste
                    self.profil_combo.clear()
                    self.profil_combo.addItem("-- Sélectionner un profil --")
                    self.load_profiles()
    
    def setup_clauses(self):
        """Configure les clauses optionnelles"""
        clauses_config = [
            {
                'name': 'Garantie d\'actif/passif (rachat de parts)',
                'text': 'Le cas échéant, production de la garantie d\'actif/passif fournie par les vendeurs et examen favorable de LCL',
                'fields': [
                    {'name': 'nom_vendeur', 'label': 'Nom du vendeur (optionnel)', 'type': 'text'}
                ]
            },
            {
                'name': 'Niveau de commercialisation (lots T3, T4, T5)',
                'text': 'Justification d\'un niveau de commercialisation incluant au moins [nombre_t3] lots de type T3 ainsi qu\'au moins [nombre_t4] lots de type T4 et [nombre_t5] lots de type T5 (attestation du Notaire indiquant le niveau de pré commercialisation)',
                'fields': [
                    {'name': 'nombre_t3', 'label': 'Nombre de lots T3', 'type': 'number'},
                    {'name': 'nombre_t4', 'label': 'Nombre de lots T4', 'type': 'number'},
                    {'name': 'nombre_t5', 'label': 'Nombre de lots T5', 'type': 'number'}
                ]
            },
            {
                'name': 'Accord de financement des réservataires',
                'text': 'Justification de l\'obtention d\'un accord de principe de financement par la majorité des réservataires',
                'fields': []
            },
            {
                'name': 'Agrément bailleur social',
                'text': 'Justification de l\'obtention de l\'agrément par [nom_bailleur_agrement] pour la partie « [type_bloc] »',
                'fields': [
                    {'name': 'nom_bailleur_agrement', 'label': 'Nom du bailleur', 'type': 'text'},
                    {'name': 'type_bloc', 'label': 'Type de bloc (social/LLS/LLI/ULS)', 'type': 'text'}
                ]
            },
            {
                'name': 'Engagement d\'information modification PC',
                'text': 'Engagement de l\'emprunteur d\'informer la banque de toute demande de PC modificatif et ce jusqu\'au remboursement complet des concours accordés',
                'fields': []
            },

            {
                'name': 'Contrat de réservation bailleur',
                'text': 'Justification d\'un contrat de réservation signé de [nom du bailleur] pour la partie « [bloc social / LLS (logements locatifs sociaux) / LLI (logements locatifs intermédiaires) / ULS (usufruit locatif social)] » comprenant nom, adresse, prix de vente TTC et échéancier des versements',
                'fields': [
                    {'name': 'nom_bailleur_reservation', 'label': 'Nom du bailleur (réservation)', 'type': 'text'},
                    {'name': 'type_bloc_reservation', 'label': 'Type de bloc (social/LLS/LLI/ULS)', 'type': 'text'}
                ]
            },
            {
                'name': 'Niveau de commercialisation libre',
                'text': "Justification d'un niveau de commercialisation du CATTC « libre » dépassant [niveau_commercialisation_libre]% du CATTC « libre » (attestation notariée indiquant le niveau de pré commercialisation) ;",
                'fields': [
                    {'name': 'niveau_commercialisation_libre', 'label': 'Niveau commercialisation libre (%)', 'type': 'number'}
                ]
            },

        ]
        
        self.clause_widgets = []
        
        for clause_config in clauses_config:
            clause_widget = ClauseWidget(
                clause_config['name'],
                clause_config['text'],
                clause_config['fields']
            )
            self.clause_widgets.append(clause_widget)
            self.clauses_layout.addWidget(clause_widget)
            
            # Ajouter un séparateur
            line = QFrame()
            line.setFrameShape(QFrame.Shape.HLine)
            line.setFrameShadow(QFrame.Shadow.Sunken)
            self.clauses_layout.addWidget(line)
    
    def import_template(self):
        """Importe un template Word"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Sélectionner le template Word",
            "",
            "Fichiers Word (*.docx);;Tous les fichiers (*)"
        )
        
        if file_path:
            self.template_path = file_path
            self.template_label.setText(f"Template: {Path(file_path).name}")
            self.generate_button.setEnabled(True)
            self.preview_button.setEnabled(True)
    
    def get_all_values(self) -> Dict[str, str]:
        """Récupère toutes les valeurs saisies"""
        values = {}
        
        # Champs principaux
        for field_key, widget in self.fields.items():
            text = widget.text().strip()
            # Pour les montants, applique le formatage à points
            if field_key in ['montant_credit', 'montant_gfa', 'frais_dossier', 'montant_apports']:
                text = format_number_with_dots(text)
            values[field_key] = text

        
        # Civilité
        values['civilite'] = self.civilite_combo.currentText()
        
        # Objet/Description
        values['objet'] = self.objet_text.toPlainText().strip()
        
        # Nouveaux champs de taux
        values['taux_speculatif'] = f"{self.taux_speculatif.value():.2f}".replace('.', ',')
        values['taux_non_speculatif'] = f"{self.taux_non_speculatif.value():.2f}".replace('.', ',')
        values['taux_comission_engagement_speculatif'] = f"{self.taux_comission_engagement_speculatif.value():.2f}".replace('.', ',')
        values['taux_comission_engagement_non_speculatif'] = f"{self.taux_comission_engagement_non_speculatif.value():.2f}".replace('.', ',')
        values['taux_comission_forfaitaire'] = f"{self.taux_comission_forfaitaire.value():.2f}".replace('.', ',')
        
        # Niveaux de commercialisation
        values['niveau_commercialisation'] = f"{int(self.niveau_commercialisation.value())}"
        
        # État des cases à cocher
        values['inclure_apports'] = self.inclure_apports_checkbox.isChecked()
        values['conditions_speculatives'] = self.conditions_speculatives_checkbox.isChecked()
        values['conditions_non_speculatives'] = self.conditions_non_speculatives_checkbox.isChecked()
        
        # Récupérer les valeurs des clauses optionnelles
        for clause_widget in self.clause_widgets:
            if clause_widget.is_enabled():
                field_values = clause_widget.get_field_values()
                values.update(field_values)
        
        # Convertir les montants en lettres
        try:
            if values['montant_credit']:
                montant = int(values['montant_credit'].replace(' ', '').replace('.', '').replace(',', ''))
                values['montant_credit_lettres'] = NumberToWords.convert(montant)

            if values['montant_gfa']:
                montant = int(values['montant_gfa'].replace(' ', '').replace('.', '').replace(',', ''))
                values['montant_gfa_lettres'] = NumberToWords.convert(montant)

            if values['frais_dossier']:
                montant = int(values['frais_dossier'].replace(' ', '').replace('.', '').replace(',', ''))
                values['frais_dossier_lettres'] = NumberToWords.convert(montant)

            if values['montant_apports']:
                montant = int(values['montant_apports'].replace(' ', '').replace('.', '').replace(',', ''))
                values['montant_apports_lettres'] = NumberToWords.convert(montant)

        except ValueError:
            pass  # Ignorer les erreurs de conversion
        
        return values
    
    def preview_termsheet(self):
        """Affiche un aperçu des valeurs qui seront remplacées"""
        if not self.template_path:
            QMessageBox.warning(self, "Erreur", "Veuillez d'abord importer un template.")
            return
        
        values = self.get_all_values()
        
        # Créer le texte d'aperçu
        preview_text = "=== APERÇU DES VALEURS ===\n\n"
        
        preview_text += "CHAMPS PRINCIPAUX:\n"
        for key, value in values.items():
            if value and not key.endswith('_lettres') and key not in ['inclure_apports', 'conditions_speculatives', 'conditions_non_speculatives']:
                preview_text += f"[{key.upper()}] = {value}\n"
        
        preview_text += "\n\nOPTIONS:\n"
        if values.get('inclure_apports'):
            preview_text += "✓ Mention '(en y ajoutant les apports)' incluse\n"
        if values.get('conditions_speculatives'):
            preview_text += "✓ Conditions spéculatives incluses\n"
        if values.get('conditions_non_speculatives'):
            preview_text += "✓ Conditions non spéculatives incluses\n"
        
        preview_text += "\n\nCLAUSES OPTIONNELLES:\n"
        for clause_widget in self.clause_widgets:
            if clause_widget.is_enabled():
                preview_text += f"✓ {clause_widget.clause_name}\n"
                field_values = clause_widget.get_field_values()
                for field_name, field_value in field_values.items():
                    if field_value:
                        preview_text += f"  - {field_name}: {field_value}\n"
            else:
                preview_text += f"✗ {clause_widget.clause_name} (désactivée)\n"
        
        # Afficher dans une boîte de dialogue
        msg = QMessageBox()
        msg.setWindowTitle("Aperçu du Termsheet")
        msg.setText(preview_text)
        msg.setStandardButtons(QMessageBox.StandardButton.Ok)
        msg.exec()
    
    def generate_termsheet(self):
        """Génère le fichier Word final"""
        if not self.template_path:
            QMessageBox.warning(self, "Erreur", "Veuillez d'abord importer un template.")
            return
        
        # AJOUTER CETTE LIGNE :
        self.check_and_propose_save_profile()
        
        try:
            # Charger le document template
            doc = Document(self.template_path)
            
            # Récupérer toutes les valeurs
            values = self.get_all_values()
            
            # Remplacer les variables dans le document
            self.replace_variables_in_document(doc, values)
            
            # Traiter les clauses optionnelles
            self.process_optional_clauses(doc)
            
            # Sauvegarder le fichier généré
            output_path = self.get_output_path()
            doc.save(output_path)
            
            # Message de succès
            reply = QMessageBox.question(
                self,
                "Termsheet généré",
                f"Le termsheet a été généré avec succès !\n\nFichier: {Path(output_path).name}\n\nVoulez-vous ouvrir le dossier de destination ?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # Ouvrir le dossier selon l'OS
                import platform
                if platform.system() == "Windows":
                    os.startfile(Path(output_path).parent)
                elif platform.system() == "Darwin":  # macOS
                    os.system(f"open '{Path(output_path).parent}'")
                else:  # Linux
                    os.system(f"xdg-open '{Path(output_path).parent}'")
        
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la génération:\n{str(e)}")
    
    def replace_variables_in_document(self, doc: Document, values: Dict[str, str]):
        """Remplace toutes les variables [VAR] dans le document"""
        # Remplacer dans les paragraphes
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, values)
        
        # Remplacer dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, values)
        
        # Remplacer dans les en-têtes et pieds de page
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self.replace_in_paragraph(paragraph, values)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self.replace_in_paragraph(paragraph, values)
    
    def replace_in_paragraph(self, paragraph, values: Dict[str, str]):
        """Remplace les variables dans un paragraphe en préservant le formatage"""
        original_text = paragraph.text
        
        # Mapping des variables vers les valeurs
        replacements = {
            '[Nom du promoteur]': values.get('nom_promoteur', ''),
            '[nom]': values.get('nom_contact', ''),
            '[Adresse du promoteur]': values.get('adresse_promoteur', ''),
            '[date]': values.get('date', ''),
            '[référence dossier]': values.get('reference_dossier', ''),
            '[Monsieur/Madame/Messieurs]': values.get('civilite', ''),
            '[NOM]': values.get('nom_sccv', ''),
            '[n° siren]': values.get('numero_siren', ''),
            '[Ville]': values.get('ville_rcs', ''),
            '[nom de la SCCV]': values.get('nom_sccv', ''),
            '[objet]': values.get('objet', ''),
            '[le bailleur]': values.get('nom_bailleur_agrement', ''),
            '[nombre_credit]': values.get('montant_credit', ''),
            '[nombre_credit_lettres]': values.get('montant_credit_lettres', ''),
            '[montant_credit]': values.get('montant_credit', ''),
            '[montant_credit_lettres]': values.get('montant_credit_lettres', ''),
            '[nombre_gfa]': values.get('montant_gfa', ''),
            '[nombre_gfa_lettres]': values.get('montant_gfa_lettres', ''),
            '[nombre_apport]': values.get('montant_apports', ''),
            '[nombre_apport_lettres]': values.get('montant_apports_lettres', ''),
            '[nombre_frais_dossier]': values.get('frais_dossier', ''),
            '[nombre_frais_dossier_lettres]': values.get('frais_dossier_lettres', ''),
            '[nombre_t3]': values.get('nombre_t3', ''),
            '[nombre_t4]': values.get('nombre_t4', ''),
            '[nombre_t5]': values.get('nombre_t5', ''),
            '[taux_speculatif]': values.get('taux_speculatif', ''),
            '[taux_non_speculatif]': values.get('taux_non_speculatif', ''),
            '[taux_comission_engagement_speculatif]': values.get('taux_comission_engagement_speculatif', ''),
            '[taux_comission_engagement_non_speculatif]': values.get('taux_comission_engagement_non_speculatif', ''),
            '[taux_comission_forfaitaire]': values.get('taux_comission_forfaitaire', ''),
            '[niveau_commercialisation_libre]': values.get('niveau_commercialisation_libre', ''),
            '[nom_bailleur_agrement]': values.get('nom_bailleur_agrement', ''),
            '[type_bloc]': values.get('type_bloc', ''),
            '[date_echeance_gfa]': values.get('date_echeance_gfa', ''),
            '[nom du bailleur]': values.get('nom_bailleur_reservation', ''),
            '[nom_bailleur_reservation]': values.get('nom_bailleur_reservation', ''),
            '[type_bloc_reservation]': values.get('type_bloc_reservation', ''),
        }
        
        # Gestion du niveau de commercialisation avec/sans apports
        if values.get('inclure_apports', False):
            replacements['[niveau_commercialisation]'] = f"{values.get('niveau_commercialisation', '')}"
            replacements['[mention_apports]'] = '(en y ajoutant les apports),'
        else:
            replacements['[niveau_commercialisation]'] = f"{values.get('niveau_commercialisation', '')}"
            replacements['[mention_apports]'] = ''
        
        # Gestion des conditions spéculatives
        if values.get('conditions_speculatives', True):
            replacements['[interets_speculatifs]'] = f"Intérêts portant sur les sommes utilisées calculés sur l'EURIBOR de la durée du tirage (minimum un mois -- maximum 12 mois) majoré de {values.get('taux_speculatif', '')}% l'an, perçus d'avance le jour de la mise à disposition des fonds ;"
            replacements['[commission_speculative]'] = f"{values.get('taux_comission_engagement_speculatif', '')}% l'an, calculée sur le montant total du crédit autorisé et perçue trimestriellement et d'avance ;"
        else:
            replacements['[interets_speculatifs]'] = ''
            replacements['[commission_speculative]'] = ''
        
        # Gestion des conditions non spéculatives
        if values.get('conditions_non_speculatives', True):
            replacements['[interets_non_speculatifs]'] = f"Lorsque le montant du CA TTC des VEFA actées atteindra 40% et plus du Prix de Revient TTC, les intérêts portant sur les sommes utilisées calculés sur l'EURIBOR de la durée du tirage (minimum un mois -- maximum 12 mois) seront ramenés à {values.get('taux_non_speculatif', '')}% l'an, perçus d'avance le jour de la mise à disposition des fonds."
            replacements['[commission_non_speculative]'] = f"Lorsque le montant du CA TTC des VEFA actées atteindra 40% et plus du Prix de Revient TTC, {values.get('taux_comission_engagement_non_speculatif', '')}% l'an, calculée sur le montant total du crédit autorisé et perçue trimestriellement et d'avance."
        else:
            replacements['[interets_non_speculatifs]'] = ''
            replacements['[commission_non_speculative]'] = ''
        
        # Gestion des clauses optionnelles
        # Clause 1: Garantie d'actif/passif
        if self.clause_widgets[0].is_enabled():
            replacements['[clause_garantie_actif_passif]'] = "Le cas échéant, production de la garantie d'actif/passif fournie par les vendeurs et examen favorable de LCL ; {cas rachat de parts de société}"
        else:
            replacements['[clause_garantie_actif_passif]'] = ''
        
        # Clause 2: Niveau de commercialisation lots
        if self.clause_widgets[1].is_enabled():
            replacements['[clause_niveau_commercialisation_lots]'] = f"Justification d'un niveau de commercialisation incluant au moins {values.get('nombre_t3', '')} lots de type T3 ainsi qu'au moins {values.get('nombre_t4', '')} lots de type T4 et {values.get('nombre_t5', '')} lots de type T5 (attestation du Notaire indiquant le niveau de pré commercialisation) ;"
        else:
            replacements['[clause_niveau_commercialisation_lots]'] = ''
        
        # Clause 3: Accord de financement
        if self.clause_widgets[2].is_enabled():
            replacements['[clause_accord_financement]'] = "Justification de l'obtention d'un accord de principe de financement par la majorité des réservataires ;"
        else:
            replacements['[clause_accord_financement]'] = ''
        
        # Clause 4: Agrément bailleur
        if self.clause_widgets[3].is_enabled():
            replacements['[clause_agrement_bailleur]'] = f"Justification de l'obtention de l'agrément par {values.get('nom_bailleur_agrement', '')} pour la partie « {values.get('type_bloc', '')} » ;"
        else:
            replacements['[clause_agrement_bailleur]'] = ''
        
        # Clause 5: Engagement PC
        if self.clause_widgets[4].is_enabled():
            replacements['[clause_engagement_pc]'] = "Engagement de l'emprunteur d'informer la banque de toute demande de PC modificatif et ce jusqu'au remboursement complet des concours accordés ;"
        else:
            replacements['[clause_engagement_pc]'] = ''
        
        # Clause 6: Contrat de réservation
        if self.clause_widgets[5].is_enabled():
            replacements['[clause_contrat_reservation]'] = f"Justification d'un contrat de réservation signé de {values.get('nom_bailleur_reservation', '')} pour la partie « {values.get('type_bloc_reservation', '')} » comprenant nom, adresse, prix de vente TTC et échéancier des versements ;"
        else:
            replacements['[clause_contrat_reservation]'] = ''

        # Clause 7: Niveau de commercialisation libre
        if len(self.clause_widgets) > 6 and self.clause_widgets[6].is_enabled():
            niveau = values.get('niveau_commercialisation_libre', '')
            replacements['[clause_niveau_commercialisation_libre]'] = (
                f"Justification d'un niveau de commercialisation du CATTC « libre » dépassant {niveau}% du CATTC « libre » (attestation notariée indiquant le niveau de pré commercialisation) ;"
                if niveau else ''
            )
        else:
            replacements['[clause_niveau_commercialisation_libre]'] = ''

        
        # Effectuer les remplacements
        new_text = original_text
        for old_text, replacement in replacements.items():
            if old_text in new_text:
                new_text = new_text.replace(old_text, replacement)
        
        # Mettre à jour le paragraphe si il y a eu des changements
        if new_text != original_text:
            # Conserver le formatage du premier run
            if paragraph.runs:
                # Sauvegarder le style du premier run
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_italic = first_run.font.italic
                
                # Effacer le contenu et ajouter le nouveau texte
                paragraph.clear()
                new_run = paragraph.add_run(new_text)
                
                # Restaurer le formatage
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                if font_bold:
                    new_run.font.bold = font_bold
                if font_italic:
                    new_run.font.italic = font_italic
            else:
                paragraph.text = new_text
    
    def process_optional_clauses(self, doc: Document):
        """Traite les clauses optionnelles - plus besoin de supprimer, tout est géré dans replace_in_paragraph"""
        # Cette méthode n'est plus nécessaire car tout est géré dans replace_in_paragraph
        # mais on la garde pour compatibilité
        pass
    
    def remove_speculative_conditions(self, doc: Document):
        """Supprime les conditions spéculatives du document"""
        speculative_patterns = [
            r'Intérêts portant sur les sommes utilisées.*?majoré de.*?fonds.*?;',
            r'0,75%.*?l\'an.*?calculée sur le montant total du crédit autorisé.*?d\'avance.*?;'
        ]
        
        for pattern in speculative_patterns:
            self.remove_clause_from_document(doc, pattern)
    
    def remove_non_speculative_conditions(self, doc: Document):
        """Supprime les conditions non spéculatives du document"""
        non_speculative_patterns = [
            r'Lorsque le montant du CA TTC des VEFA actées atteindra 40%.*?fonds\.',
            r'Lorsque le montant du CA TTC des VEFA actées atteindra 40%.*?d\'avance\.'
        ]
        
        for pattern in non_speculative_patterns:
            self.remove_clause_from_document(doc, pattern)
    
    def remove_clause_from_document(self, doc: Document, pattern: str):
        """Supprime une clause du document"""
        paragraphs_to_remove = []
        
        # Rechercher dans tous les paragraphes
        for paragraph in doc.paragraphs:
            if re.search(pattern, paragraph.text, re.IGNORECASE | re.DOTALL):
                paragraphs_to_remove.append(paragraph)
        
        # Rechercher dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_paragraphs_to_remove = []
                    for paragraph in cell.paragraphs:
                        if re.search(pattern, paragraph.text, re.IGNORECASE | re.DOTALL):
                            cell_paragraphs_to_remove.append(paragraph)
                    
                    # Supprimer les paragraphes des cellules
                    for paragraph in cell_paragraphs_to_remove:
                        p = paragraph._element
                        p.getparent().remove(p)
        
        # Supprimer les paragraphes identifiés
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)
    
    def get_output_path(self) -> str:
        """Génère le chemin de sortie pour le fichier"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        template_name = Path(self.template_path).stem
        output_name = f"{template_name}_genere_{timestamp}.docx"
        
        # Demander à l'utilisateur où sauvegarder
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "Sauvegarder le termsheet généré",
            output_name,
            "Fichiers Word (*.docx);;Tous les fichiers (*)"
        )
        
        if not output_path:
            # Si l'utilisateur annule, sauvegarder dans le dossier courant
            output_path = f"./{output_name}"
        
        return output_path


def main():
    """Fonction principale"""
    # Vérifier que les dépendances sont installées
    try:
        import docx
    except ImportError:
        print("ERREUR: Le module python-docx n'est pas installé.")
        print("Veuillez l'installer avec la commande :")
        print("pip install python-docx")
        print("\nPuis relancez le programme.")
        input("Appuyez sur Entrée pour fermer...")
        return
    
    app = QApplication(sys.argv)
    app.setApplicationName("Générateur de Termsheet LCL")
    
    # Définir l'icône de l'application si disponible
    try:
        app.setWindowIcon(QPixmap("icon.png"))  # Optionnel
    except:
        pass
    
    # Style de l'application - Thème LCL
    app.setStyleSheet("""
        QMainWindow {
            background-color: white;
            color: #333333;
        }
        QWidget {
            background-color: white;
            color: #333333;
        }
        QGroupBox {
            font-weight: bold;
            color: #00468E;
            font-size: 12px;
            border: 2px solid #00468E;
            border-radius: 5px;
            margin-top: 1ex;
            padding-top: 10px;
            background-color: white;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 5px 0 5px;
            color: #00468E;
            font-weight: bold;
        }
        QPushButton {
            background-color: #00468E;
            color: white;
            border: none;
            padding: 8px 16px;
            border-radius: 4px;
            font-weight: bold;
            font-size: 11px;
        }
        QPushButton:hover {
            background-color: #003a75;
        }
        QPushButton:pressed {
            background-color: #002d5c;
        }
        QPushButton:disabled {
            background-color: #cccccc;
            color: #666666;
        }
        QLineEdit, QTextEdit, QSpinBox, QComboBox, QDoubleSpinBox {
            border: 2px solid #cccccc;
            border-radius: 3px;
            padding: 5px;
            background-color: white;
            color: #333333;
            font-weight: bold;
        }
        QLineEdit:focus, QTextEdit:focus, QSpinBox:focus, QComboBox:focus, QDoubleSpinBox:focus {
            border: 2px solid #00468E;
        }
        QLabel {
            color: #333333;
            font-weight: bold;
            background-color: transparent;
        }
        QCheckBox {
            spacing: 5px;
            color: #333333;
            font-weight: bold;
            background-color: transparent;
        }
        QCheckBox::indicator {
            width: 18px;
            height: 18px;
        }
        QCheckBox::indicator:unchecked {
            border: 2px solid #cccccc;
            background-color: white;
            border-radius: 3px;
        }
        QCheckBox::indicator:checked {
            border: 2px solid #00468E;
            background-color: #00468E;
            border-radius: 3px;
        }
        QScrollArea {
            border: none;
            background-color: white;
        }
        QScrollArea > QWidget > QWidget {
            background-color: white;
        }
        QFrame[frameShape="4"] {
            color: #cccccc;
        }
    """)
    
    window = TermsheetGenerator()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()

