#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur automatique de Termsheet CII (Caution d'Indemnité d'Immobilisation) Word
Interface PyQt6 pour remplir un template de termsheet CII et générer un fichier Word final
"""

import sys
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QFormLayout, QLineEdit, QTextEdit, QCheckBox, QLabel, QPushButton,
    QFileDialog, QMessageBox, QScrollArea, QGroupBox, QSpinBox,
    QComboBox, QFrame, QDoubleSpinBox, QDateEdit
)
from PyQt6.QtCore import Qt, pyqtSignal, QDate
from PyQt6.QtGui import QFont, QPixmap

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Erreur: Le module python-docx n'est pas installé.")
    print("Installez-le avec: pip install python-docx")
    sys.exit(1)


class NumberToWords:
    """Convertit les nombres en lettres (français)"""
    
    UNITS = ['', 'un', 'deux', 'trois', 'quatre', 'cinq', 'six', 'sept', 'huit', 'neuf']
    TEENS = ['dix', 'onze', 'douze', 'treize', 'quatorze', 'quinze', 'seize', 
             'dix-sept', 'dix-huit', 'dix-neuf']
    TENS = ['', '', 'vingt', 'trente', 'quarante', 'cinquante', 'soixante', 
            'soixante-dix', 'quatre-vingt', 'quatre-vingt-dix']
    
    @classmethod
    def convert(cls, number: int) -> str:
        """Convertit un nombre en lettres"""
        if number == 0:
            return "zéro"
        
        if number < 0:
            return "moins " + cls.convert(-number)
        
        result = []
        
        # Milliards
        if number >= 1000000000:
            billions = number // 1000000000
            if billions == 1:
                result.append("un milliard")
            else:
                result.append(cls._convert_hundreds(billions) + " milliards")
            number %= 1000000000
        
        # Millions
        if number >= 1000000:
            millions = number // 1000000
            if millions == 1:
                result.append("un million")
            else:
                result.append(cls._convert_hundreds(millions) + " millions")
            number %= 1000000
        
        # Milliers
        if number >= 1000:
            thousands = number // 1000
            if thousands == 1:
                result.append("mille")
            else:
                result.append(cls._convert_hundreds(thousands) + " mille")
            number %= 1000
        
        # Centaines
        if number > 0:
            result.append(cls._convert_hundreds(number))
        
        return " ".join(result)
    
    @classmethod
    def _convert_hundreds(cls, number: int) -> str:
        """Convertit un nombre de 0 à 999 en lettres"""
        result = []
        
        # Centaines
        if number >= 100:
            hundreds = number // 100
            if hundreds == 1:
                result.append("cent")
            else:
                result.append(cls.UNITS[hundreds] + " cent")
            number %= 100
        
        # Dizaines et unités
        if number >= 20:
            tens = number // 10
            units = number % 10
            if tens == 7:
                if units == 1:
                    result.append("soixante et onze")
                elif units > 1:
                    result.append("soixante-" + cls.TEENS[units - 10])
                else:
                    result.append("soixante-dix")
            elif tens == 9:
                if units > 0:
                    result.append("quatre-vingt-" + cls.TEENS[units - 10])
                else:
                    result.append("quatre-vingt-dix")
            else:
                tens_word = cls.TENS[tens]
                if units == 1 and tens != 8:
                    result.append(tens_word + " et un")
                elif units > 0:
                    result.append(tens_word + "-" + cls.UNITS[units])
                else:
                    if tens == 8:
                        result.append("quatre-vingts")
                    else:
                        result.append(tens_word)
        elif number >= 10:
            result.append(cls.TEENS[number - 10])
        elif number > 0:
            result.append(cls.UNITS[number])
        
        return " ".join(result)
    

def format_number_with_dots(number: str) -> str:
    """Formatte un nombre sous forme 1.000.000"""
    try:
        n = int(number.replace(" ", "").replace(",", "").replace(".", ""))
        return f"{n:,}".replace(",", ".")
    except:
        return number


class TermsheetCIIGenerator(QMainWindow):
    """Application principale pour générer les termsheets CII"""
    
    def __init__(self):
        super().__init__()
        self.template_path = None
        self.setup_ui()
        self.load_default_template()
    
    def load_default_template(self):
        """Charge automatiquement le template par défaut"""
        default_template = "template_cii.docx"
        
        if os.path.exists(default_template):
            self.template_path = default_template
            self.template_label.setText(f"Template: {default_template}")
            self.generate_button.setEnabled(True)
            self.preview_button.setEnabled(True)
        else:
            self.template_label.setText("Template par défaut non trouvé - Veuillez importer un template")
            self.generate_button.setEnabled(False)
            self.preview_button.setEnabled(False)
    
    def setup_ui(self):
        """Configure l'interface utilisateur"""
        self.setWindowTitle("Générateur de Termsheet CII - LCL")
        self.setGeometry(100, 100, 900, 700)
        
        # Widget principal avec scroll
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        main_layout = QVBoxLayout()
        
        # Zone de scroll
        scroll = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout()
        
        # Section import template
        template_group = QGroupBox("Template Word")
        template_layout = QHBoxLayout()
        
        self.template_label = QLabel("Chargement du template...")
        self.template_button = QPushButton("Changer de Template (.docx)")
        self.template_button.clicked.connect(self.import_template)
        
        template_layout.addWidget(self.template_label)
        template_layout.addWidget(self.template_button)
        template_group.setLayout(template_layout)
        
        scroll_layout.addWidget(template_group)
        
        # Section informations principales
        main_info_group = QGroupBox("Informations Principales")
        main_info_layout = QFormLayout()
        
        # Créer tous les champs principaux
        self.fields = {}
        
        main_fields = [
            ('nom_promoteur', 'Nom du promoteur'),
            ('nom_contact', 'Nom du contact'),
            ('adresse_promoteur', 'Adresse du promoteur'),
            ('date', 'Date (format DD/MM/YYYY)'),
            ('reference_dossier', 'Référence dossier'),
            ('nom_sccv', 'Nom de la SCCV (Emprunteur)'),
            ('numero_siren', 'Numéro SIREN'),
            ('ville_rcs', 'Ville RCS'),
        ]
        
        for field_key, field_label in main_fields:
            widget = QLineEdit()
            self.fields[field_key] = widget
            main_info_layout.addRow(field_label, widget)
        
        # Civilité (dropdown)
        self.civilite_combo = QComboBox()
        self.civilite_combo.addItems(['Monsieur', 'Madame', 'Messieurs'])
        main_info_layout.addRow('Civilité', self.civilite_combo)
        
        # Objet/Description (zone de texte plus grande)
        self.objet_text = QTextEdit()
        self.objet_text.setMaximumHeight(80)
        self.objet_text.setPlainText("")
        main_info_layout.addRow('Objet / Description du programme', self.objet_text)
        
        main_info_group.setLayout(main_info_layout)
        scroll_layout.addWidget(main_info_group)
        
        # Section CII (Caution d'Indemnité d'Immobilisation)
        cii_group = QGroupBox("Cautions d'Indemnité d'Immobilisation (CII)")
        cii_main_layout = QVBoxLayout()
        
        # Container pour les CII
        self.cii_container = QWidget()
        self.cii_layout = QVBoxLayout()
        self.cii_container.setLayout(self.cii_layout)
        
        # Liste pour stocker les widgets CII
        self.cii_widgets = []
        
        # Boutons pour gérer les CII
        cii_buttons_layout = QHBoxLayout()
        self.add_cii_button = QPushButton("Ajouter une CII")
        self.add_cii_button.clicked.connect(self.add_cii)
        self.remove_cii_button = QPushButton("Supprimer la dernière CII")
        self.remove_cii_button.clicked.connect(self.remove_cii)
        
        cii_buttons_layout.addWidget(self.add_cii_button)
        cii_buttons_layout.addWidget(self.remove_cii_button)
        cii_buttons_layout.addStretch()
        
        cii_main_layout.addLayout(cii_buttons_layout)
        cii_main_layout.addWidget(self.cii_container)
        
        cii_group.setLayout(cii_main_layout)
        scroll_layout.addWidget(cii_group)
        
        # Ajouter une CII par défaut
        self.add_cii()
        
        # Section Conditions de rémunération
        remun_group = QGroupBox("Conditions de Rémunération")
        remun_layout = QFormLayout()
        
        # Commission forfaitaire
        self.commission_forfaitaire = QLineEdit()
        self.commission_forfaitaire.setPlaceholderText("Montant en euros")
        remun_layout.addRow('Commission forfaitaire globale (€)', self.commission_forfaitaire)
        
        # Taux commission de risque
        self.taux_commission_risque = QDoubleSpinBox()
        self.taux_commission_risque.setDecimals(2)
        self.taux_commission_risque.setMaximum(99.99)
        self.taux_commission_risque.setValue(0.50)
        self.taux_commission_risque.setSuffix(' %')
        remun_layout.addRow('Taux commission de risque', self.taux_commission_risque)
        
        # Frais d'acte
        self.frais_acte = QLineEdit()
        self.frais_acte.setPlaceholderText("Montant en euros")
        self.frais_acte.setText("290")
        remun_layout.addRow('Frais d\'acte (€)', self.frais_acte)
        
        remun_group.setLayout(remun_layout)
        scroll_layout.addWidget(remun_group)
        
        # Section Modalités
        modalites_group = QGroupBox("Modalités")
        modalites_layout = QFormLayout()
        
        # Commission de retainer
        self.commission_retainer = QLineEdit()
        self.commission_retainer.setPlaceholderText("Montant en euros")
        modalites_layout.addRow('Commission de retainer (€)', self.commission_retainer)
        
        # Date de validité de l'accord
        self.date_validite_accord = QLineEdit()
        self.date_validite_accord.setPlaceholderText("Ex: 22 juin 2025")
        modalites_layout.addRow('Date de fin de validité de l\'accord', self.date_validite_accord)
        
        modalites_group.setLayout(modalites_layout)
        scroll_layout.addWidget(modalites_group)
        
        # Boutons d'action
        buttons_layout = QHBoxLayout()
        
        self.generate_button = QPushButton("Générer le Termsheet CII")
        self.generate_button.clicked.connect(self.generate_termsheet)
        self.generate_button.setEnabled(False)
        
        self.preview_button = QPushButton("Aperçu")
        self.preview_button.clicked.connect(self.preview_termsheet)
        self.preview_button.setEnabled(False)
        
        buttons_layout.addWidget(self.preview_button)
        buttons_layout.addWidget(self.generate_button)
        
        scroll_layout.addLayout(buttons_layout)
        
        scroll_widget.setLayout(scroll_layout)
        scroll.setWidget(scroll_widget)
        scroll.setWidgetResizable(True)
        
        main_layout.addWidget(scroll)
        main_widget.setLayout(main_layout)
    
    def add_cii(self):
        """Ajoute une nouvelle CII"""
        cii_index = len(self.cii_widgets)
        
        # Créer le widget pour cette CII
        cii_widget = QGroupBox(f"CII #{cii_index + 1}")
        cii_form_layout = QFormLayout()
        
        # Dictionnaire pour stocker les champs de cette CII
        cii_fields = {}
        
        # Bénéficiaires
        beneficiaires = QLineEdit()
        beneficiaires.setPlaceholderText("Ex: Madame Marie DUPONT et Monsieur Pierre MARTIN")
        cii_fields['beneficiaires'] = beneficiaires
        cii_form_layout.addRow('Bénéficiaires (en faveur de)', beneficiaires)
        
        # Venant au droit de (optionnel)
        venant_au_droit = QLineEdit()
        venant_au_droit.setPlaceholderText("Optionnel - Ex: Monsieur Jean MARTIN")
        cii_fields['venant_au_droit'] = venant_au_droit
        cii_form_layout.addRow('Venant au droit de (optionnel)', venant_au_droit)
        
        # Montant
        montant = QLineEdit()
        montant.setPlaceholderText("Montant en euros")
        cii_fields['montant'] = montant
        cii_form_layout.addRow('Montant (€)', montant)
        
        # Date d'échéance
        date_echeance = QLineEdit()
        date_echeance.setPlaceholderText("Ex: 31 juillet 2025")
        cii_fields['date_echeance'] = date_echeance
        cii_form_layout.addRow('Date d\'échéance', date_echeance)
        
        cii_widget.setLayout(cii_form_layout)
        
        # Stocker le widget et ses champs
        self.cii_widgets.append({
            'widget': cii_widget,
            'fields': cii_fields
        })
        
        # Ajouter à l'interface
        self.cii_layout.addWidget(cii_widget)
        
        # Mettre à jour l'état du bouton supprimer
        self.remove_cii_button.setEnabled(len(self.cii_widgets) > 1)
    
    def remove_cii(self):
        """Supprime la dernière CII"""
        if len(self.cii_widgets) > 1:
            # Récupérer et supprimer le dernier widget
            last_cii = self.cii_widgets.pop()
            last_cii['widget'].setParent(None)
            last_cii['widget'].deleteLater()
            
            # Mettre à jour l'état du bouton supprimer
            self.remove_cii_button.setEnabled(len(self.cii_widgets) > 1)
            
            # Mettre à jour les titres des CII restantes
            for i, cii_data in enumerate(self.cii_widgets):
                cii_data['widget'].setTitle(f"CII #{i + 1}")
    
    def import_template(self):
        """Importe un template Word"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Sélectionner le template Word",
            "",
            "Fichiers Word (*.docx);;Tous les fichiers (*)"
        )
        
        if file_path:
            self.template_path = file_path
            self.template_label.setText(f"Template: {Path(file_path).name}")
            self.generate_button.setEnabled(True)
            self.preview_button.setEnabled(True)
    
    def get_all_values(self) -> Dict[str, str]:
        """Récupère toutes les valeurs saisies"""
        values = {}
        
        # Champs principaux
        for field_key, widget in self.fields.items():
            values[field_key] = widget.text().strip()
        
        # Civilité
        values['civilite'] = self.civilite_combo.currentText()
        
        # Objet/Description
        values['objet'] = self.objet_text.toPlainText().strip()
        
        # Générer la section complète des CII
        values['section_complete_cii'] = self.generate_cii_section()
        
        # Montants autres (pour compatibilité si besoin)
        values['commission_forfaitaire'] = format_number_with_dots(self.commission_forfaitaire.text().strip())
        values['frais_acte'] = format_number_with_dots(self.frais_acte.text().strip())
        values['commission_retainer'] = format_number_with_dots(self.commission_retainer.text().strip())
        
        # Dates
        values['date_validite_accord'] = self.date_validite_accord.text().strip()
        
        # Taux
        values['taux_commission_risque'] = f"{self.taux_commission_risque.value():.2f}".replace('.', ',')
        
        # Convertir les montants en lettres
        try:
            if values['commission_forfaitaire']:
                montant = int(values['commission_forfaitaire'].replace(' ', '').replace('.', '').replace(',', ''))
                values['commission_forfaitaire_lettres'] = NumberToWords.convert(montant)
            
            if values['frais_acte']:
                montant = int(values['frais_acte'].replace(' ', '').replace('.', '').replace(',', ''))
                values['frais_acte_lettres'] = NumberToWords.convert(montant)
            
            if values['commission_retainer']:
                montant = int(values['commission_retainer'].replace(' ', '').replace('.', '').replace(',', ''))
                values['commission_retainer_lettres'] = NumberToWords.convert(montant)
                
        except ValueError:
            pass  # Ignorer les erreurs de conversion
        
        return values
    
    def generate_cii_section(self) -> str:
        """Génère la section complète des CII"""
        if not self.cii_widgets:
            return ""
        
        cii_sections = []
        
        for i, cii_data in enumerate(self.cii_widgets):
            fields = cii_data['fields']
            
            # Récupérer les valeurs
            beneficiaires = fields['beneficiaires'].text().strip()
            venant_au_droit = fields['venant_au_droit'].text().strip()
            montant_str = fields['montant'].text().strip()
            date_echeance = fields['date_echeance'].text().strip()
            
            if not beneficiaires or not montant_str or not date_echeance:
                continue  # Ignorer les CII incomplètes
            
            # Formater le montant
            montant_formate = format_number_with_dots(montant_str)
            
            # Convertir en lettres
            try:
                montant_int = int(montant_str.replace(' ', '').replace(',', '').replace('.', ''))
                montant_lettres = NumberToWords.convert(montant_int)
            except ValueError:
                montant_lettres = ""
            
            # Construire le texte de la CII
            cii_text = "Caution d'indemnité d'immobilisation (CII) :\n\n"
            
            # Point a. Nature avec bénéficiaires
            nature_text = f"a. Caution d'indemnité d'immobilisation (CII), émise en faveur de {beneficiaires}"
            if venant_au_droit:
                nature_text += f", venant au droit de {venant_au_droit}"
            nature_text += ".\n\n"
            
            cii_text += nature_text
            
            # Point b. Montant
            cii_text += f"b. Montant : {montant_formate} €"
            if montant_lettres:
                cii_text += f" ({montant_lettres} euros)"
            cii_text += ".\n\n"
            
            # Point c. Date d'échéance
            cii_text += f"c. Date d'échéance : {date_echeance}.\n\n"
            
            cii_sections.append(cii_text)
        
        return "\n".join(cii_sections)
    
    def preview_termsheet(self):
        """Affiche un aperçu des valeurs qui seront remplacées"""
        if not self.template_path:
            QMessageBox.warning(self, "Erreur", "Veuillez d'abord importer un template.")
            return
        
        values = self.get_all_values()
        
        # Créer le texte d'aperçu
        preview_text = "=== APERÇU DES VALEURS CII ===\n\n"
        
        preview_text += "INFORMATIONS GÉNÉRALES:\n"
        for key in ['nom_promoteur', 'nom_contact', 'adresse_promoteur', 'date', 'reference_dossier', 'civilite', 'nom_sccv', 'numero_siren', 'ville_rcs']:
            if values.get(key):
                preview_text += f"[{key.upper()}] = {values[key]}\n"
        
        preview_text += f"\nOBJET:\n{values.get('objet', '')}\n"
        
        preview_text += "\nCAUTIONS D'INDEMNITÉ D'IMMOBILISATION:\n"
        preview_text += f"Nombre de CII: {len(self.cii_widgets)}\n"
        
        for i, cii_data in enumerate(self.cii_widgets):
            fields = cii_data['fields']
            beneficiaires = fields['beneficiaires'].text().strip()
            venant_au_droit = fields['venant_au_droit'].text().strip()
            montant = fields['montant'].text().strip()
            date_echeance = fields['date_echeance'].text().strip()
            
            preview_text += f"\nCII #{i+1}:\n"
            if beneficiaires:
                preview_text += f"  - Bénéficiaires: {beneficiaires}\n"
            if venant_au_droit:
                preview_text += f"  - Venant au droit de: {venant_au_droit}\n"
            if montant:
                montant_formate = format_number_with_dots(montant)
                try:
                    montant_int = int(montant.replace(' ', '').replace(',', '').replace('.', ''))
                    montant_lettres = NumberToWords.convert(montant_int)
                    preview_text += f"  - Montant: {montant_formate} € ({montant_lettres} euros)\n"
                except:
                    preview_text += f"  - Montant: {montant_formate} €\n"
            if date_echeance:
                preview_text += f"  - Date d'échéance: {date_echeance}\n"
        
        preview_text += "\nCONDITIONS DE RÉMUNÉRATION:\n"
        if values.get('commission_forfaitaire'):
            preview_text += f"Commission forfaitaire: {values['commission_forfaitaire']} € ({values.get('commission_forfaitaire_lettres', '')})\n"
        preview_text += f"Taux commission de risque: {values.get('taux_commission_risque', '')}%\n"
        if values.get('frais_acte'):
            preview_text += f"Frais d'acte: {values['frais_acte']} € ({values.get('frais_acte_lettres', '')})\n"
        
        preview_text += "\nMODALITÉS:\n"
        if values.get('commission_retainer'):
            preview_text += f"Commission de retainer: {values['commission_retainer']} € ({values.get('commission_retainer_lettres', '')})\n"
        if values.get('date_validite_accord'):
            preview_text += f"Date de validité: {values['date_validite_accord']}\n"
        
        # Afficher dans une boîte de dialogue
        msg = QMessageBox()
        msg.setWindowTitle("Aperçu du Termsheet CII")
        msg.setText(preview_text)
        msg.setStandardButtons(QMessageBox.StandardButton.Ok)
        msg.exec()
    
    def generate_termsheet(self):
        """Génère le fichier Word final"""
        if not self.template_path:
            QMessageBox.warning(self, "Erreur", "Veuillez d'abord importer un template.")
            return
        
        try:
            # Charger le document template
            doc = Document(self.template_path)
            
            # Récupérer toutes les valeurs
            values = self.get_all_values()
            
            # Remplacer les variables dans le document
            self.replace_variables_in_document(doc, values)
            
            # Sauvegarder le fichier généré
            output_path = self.get_output_path()
            doc.save(output_path)
            
            # Message de succès
            reply = QMessageBox.question(
                self,
                "Termsheet CII généré",
                f"Le termsheet CII a été généré avec succès !\n\nFichier: {Path(output_path).name}\n\nVoulez-vous ouvrir le dossier de destination ?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # Ouvrir le dossier selon l'OS
                import platform
                if platform.system() == "Windows":
                    os.startfile(Path(output_path).parent)
                elif platform.system() == "Darwin":  # macOS
                    os.system(f"open '{Path(output_path).parent}'")
                else:  # Linux
                    os.system(f"xdg-open '{Path(output_path).parent}'")
        
        except Exception as e:
            QMessageBox.critical(self, "Erreur", f"Erreur lors de la génération:\n{str(e)}")
    
    def replace_variables_in_document(self, doc: Document, values: Dict[str, str]):
        """Remplace toutes les variables [VAR] dans le document"""
        # Remplacer dans les paragraphes
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, values)
        
        # Remplacer dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, values)
        
        # Remplacer dans les en-têtes et pieds de page
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self.replace_in_paragraph(paragraph, values)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self.replace_in_paragraph(paragraph, values)
    
    def replace_in_paragraph(self, paragraph, values: Dict[str, str]):
        """Remplace les variables dans un paragraphe en préservant le formatage"""
        original_text = paragraph.text
        
        # Mapping des variables vers les valeurs
        replacements = {
            '[Nom du promoteur]': values.get('nom_promoteur', ''),
            '[nom]': values.get('nom_contact', ''),
            '[Adresse du promoteur]': values.get('adresse_promoteur', ''),
            '[date]': values.get('date', ''),
            '[réference dossier]': values.get('reference_dossier', ''),
            '[Monsieur/Madame/Messieurs]': values.get('civilite', ''),
            '[NOM]': values.get('nom_sccv', ''),
            '[n° siren]': values.get('numero_siren', ''),
            '[Ville]': values.get('ville_rcs', ''),
            '[objet]': values.get('objet', ''),
            '[section_complete_cii]': values.get('section_complete_cii', ''),
            '[nombre_comission_forfaitaire]': values.get('commission_forfaitaire', ''),
            '[nombre_comission_forfaitaire_lettres]': values.get('commission_forfaitaire_lettres', ''),
            '[taux_commission_risque]': values.get('taux_commission_risque', ''),
            '[nombre_frais_acte]': values.get('frais_acte', ''),
            '[nombre_frais_acte_lettres]': values.get('frais_acte_lettres', ''),
            '[nombre_commission_retainer]': values.get('commission_retainer', ''),
            '[nombre_commission_retainer_lettres]': values.get('commission_retainer_lettres', ''),
            '[date_validite_accord]': values.get('date_validite_accord', ''),
        }
        
        # Effectuer les remplacements
        new_text = original_text
        for old_text, replacement in replacements.items():
            if old_text in new_text:
                new_text = new_text.replace(old_text, replacement)
        
        # Mettre à jour le paragraphe si il y a eu des changements
        if new_text != original_text:
            # Conserver le formatage du premier run
            if paragraph.runs:
                # Sauvegarder le style du premier run
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_italic = first_run.font.italic
                
                # Effacer le contenu et ajouter le nouveau texte
                paragraph.clear()
                new_run = paragraph.add_run(new_text)
                
                # Restaurer le formatage
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                if font_bold:
                    new_run.font.bold = font_bold
                if font_italic:
                    new_run.font.italic = font_italic
            else:
                paragraph.text = new_text
    
    def get_output_path(self) -> str:
        """Génère le chemin de sortie pour le fichier"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        template_name = Path(self.template_path).stem
        output_name = f"{template_name}_CII_genere_{timestamp}.docx"
        
        # Demander à l'utilisateur où sauvegarder
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "Sauvegarder le termsheet CII généré",
            output_name,
            "Fichiers Word (*.docx);;Tous les fichiers (*)"
        )
        
        if not output_path:
            # Si l'utilisateur annule, sauvegarder dans le dossier courant
            output_path = f"./{output_name}"
        
        return output_path


def main():
    """Fonction principale"""
    # Vérifier que les dépendances sont installées
    try:
        import docx
    except ImportError:
        print("ERREUR: Le module python-docx n'est pas installé.")
        print("Veuillez l'installer avec la commande :")
        print("pip install python-docx")
        print("\nPuis relancez le programme.")
        input("Appuyez sur Entrée pour fermer...")
        return
    
    app = QApplication(sys.argv)
    app.setApplicationName("Générateur de Termsheet CII - LCL")
    
    # Définir l'icône de l'application si disponible
    try:
        app.setWindowIcon(QPixmap("icon.png"))  # Optionnel
    except:
        pass
    
    # Style de l'application - Thème LCL
    app.setStyleSheet("""
        QMainWindow {
            background-color: white;
            color: #333333;
        }
        QWidget {
            background-color: white;
            color: #333333;
        }
        QGroupBox {
            font-weight: bold;
            color: #00468E;
            font-size: 12px;
            border: 2px solid #00468E;
            border-radius: 5px;
            margin-top: 1ex;
            padding-top: 10px;
            background-color: white;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 5px 0 5px;
            color: #00468E;
            font-weight: bold;
        }
        QPushButton {
            background-color: #00468E;
            color: white;
            border: none;
            padding: 8px 16px;
            border-radius: 4px;
            font-weight: bold;
            font-size: 11px;
        }
        QPushButton:hover {
            background-color: #003a75;
        }
        QPushButton:pressed {
            background-color: #002d5c;
        }
        QPushButton:disabled {
            background-color: #cccccc;
            color: #666666;
        }
        QLineEdit, QTextEdit, QSpinBox, QComboBox, QDoubleSpinBox {
            border: 2px solid #cccccc;
            border-radius: 3px;
            padding: 5px;
            background-color: white;
            color: #333333;
            font-weight: bold;
        }
        QLineEdit:focus, QTextEdit:focus, QSpinBox:focus, QComboBox:focus, QDoubleSpinBox:focus {
            border: 2px solid #00468E;
        }
        QLabel {
            color: #333333;
            font-weight: bold;
            background-color: transparent;
        }
        QCheckBox {
            spacing: 5px;
            color: #333333;
            font-weight: bold;
            background-color: transparent;
        }
        QCheckBox::indicator {
            width: 18px;
            height: 18px;
        }
        QCheckBox::indicator:unchecked {
            border: 2px solid #cccccc;
            background-color: white;
            border-radius: 3px;
        }
        QCheckBox::indicator:checked {
            border: 2px solid #00468E;
            background-color: #00468E;
            border-radius: 3px;
        }
        QScrollArea {
            border: none;
            background-color: white;
        }
        QScrollArea > QWidget > QWidget {
            background-color: white;
        }
        QFrame[frameShape="4"] {
            color: #cccccc;
        }
    """)
    
    window = TermsheetCIIGenerator()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
